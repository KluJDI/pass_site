from flask import Flask, request, send_file, render_template
from docx import Document
import io
import os
import logging
import re

app = Flask(__name__)

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@app.route('/')
def form():
    logger.info("Открыта страница формы")
    try:
        return render_template('index.html')
    except Exception as e:
        logger.error(f"Ошибка при рендеринге формы: {str(e)}")
        return "Ошибка загрузки формы", 500

@app.route('/generate', methods=['POST'])
def generate():
    try:
        # Путь к шаблону
        template_path = os.path.join(os.path.dirname(__file__), "pasport_mesta_massovogo_prebuvaniya_lyudei.docx")
        logger.info(f"Проверка шаблона по пути: {template_path}")

        if not os.path.exists(template_path):
            logger.error(f"Шаблон {template_path} не найден")
            return "Ошибка: Шаблон документа не найден", 500

        # Текстовые поля с уникальными плейсхолдерами
        text_fields = {
            'security_classification': '[security_classification]',
            'copy_number': '[copy_number]',
            'executive_authority_head': '[executive_authority_head]',
            'executive_authority_name': '[executive_authority_name]',
            'approval_date': '[approval_date]',
            'security_agency_head': '[security_agency_head]',
            'security_agency_name': '[security_agency_name]',
            'security_agency_date': '[security_agency_date]',
            'mvd_head': '[mvd_head]',
            'mvd_name': '[mvd_name]',
            'mvd_date': '[mvd_date]',
            'mchs_head': '[mchs_head]',
            'mchs_name': '[mchs_name]',
            'mchs_date': '[mchs_date]',
            'rosgvardia_head': '[rosgvardia_head]',
            'rosgvardia_name': '[rosgvardia_name]',
            'rosgvardia_date': '[rosgvardia_date]',
            'locality': '[locality]',
            'object_name': '[object_name]',
            'object_address': '[object_address]',
            'object_affiliation': '[object_affiliation]',
            'object_boundaries': '[object_boundaries]',
            'object_area_perimeter': '[object_area_perimeter]',
            'monitoring_results': '[monitoring_results]',
            'object_category': '[object_category]',
            'mvd_territory': '[mvd_territory]',
            'public_organizations': '[public_organizations]',
            'terrain_characteristics': '[terrain_characteristics]',
            'staff_count': '[staff_count]',
            'attendance': '[attendance]',
            'tenants_info': '[tenants_info]',
            'illegal_actions_a': '[illegal_actions_a]',
            'diversion_manifestations_b': '[diversion_manifestations_b]',
            'security_forces_a': '[security_forces_a]',
            'patrol_routes_b': '[patrol_routes_b]',
            'stationary_posts_b': '[stationary_posts_b]',
            'public_guards_d': '[public_guards_d]',
            'security_equipment_e': '[security_equipment_e]',
            'notification_system_zh': '[notification_system_zh]',
            'notification_system_zh_2': '[notification_system_zh_2]',
            'notification_system_zh_3': '[notification_system_zh_3]',
            'notification_system_zh_4': '[notification_system_zh_4]',
            'notification_system_zh_5': '[notification_system_zh_5]',
            'notification_system_zh_6': '[notification_system_zh_6]',
            'technical_security_a': '[technical_security_a]',
            'fire_safety_b': '[fire_safety_b]',
            'evacuation_system_v': '[evacuation_system_v]',
            'security_reliability_a': '[security_reliability_a]',
            'urgent_measures_b': '[urgent_measures_b]',
            'funding_v': '[funding_v]',
            'additional_info': '[additional_info]',
            'recreation_areas': '[recreation_areas]',
            'communication_schemes': '[communication_schemes]',
            'evacuation_instructions': '[evacuation_instructions]',
            'correction_log': '[correction_log]',
            'rights_holder': '[rights_holder]',
            'rights_holder_name': '[rights_holder_name]',
            'creation_date': '[creation_date]',
            'update_date': '[update_date]'
        }

        # Собираем данные из формы
        fields = {key: request.form.get(key, '').strip() for key in text_fields}

        # Таблицы
        table_fields = {
            'objects_on_territory': {
                'num': request.form.getlist('object_on_territory_num[]'),
                'name': request.form.getlist('object_on_territory_name[]'),
                'details': request.form.getlist('object_on_territory_details[]'),
                'location': request.form.getlist('object_on_territory_location[]'),
                'security': request.form.getlist('object_on_territory_security[]')
            },
            'objects_nearby': {
                'num': request.form.getlist('object_nearby_num[]'),
                'name': request.form.getlist('object_nearby_name[]'),
                'characteristics': request.form.getlist('object_nearby_characteristics[]'),
                'location': request.form.getlist('object_nearby_location[]'),
                'distance': request.form.getlist('object_nearby_distance[]')
            },
            'transport_communications': {
                'num': request.form.getlist('transport_num[]'),
                'type': request.form.getlist('transport_type[]'),
                'name': request.form.getlist('transport_name[]'),
                'distance': request.form.getlist('transport_distance[]')
            },
            'service_organizations': {
                'num': request.form.getlist('service_org_num[]'),
                'name': request.form.getlist('service_org_name[]'),
                'activity': request.form.getlist('service_org_activity[]'),
                'schedule': request.form.getlist('service_org_schedule[]')
            },
            'dangerous_areas': {
                'num': request.form.getlist('dangerous_area_num[]'),
                'name': request.form.getlist('dangerous_area_name[]'),
                'worker_count': request.form.getlist('dangerous_area_worker_count[]'),
                'emergency_type': request.form.getlist('dangerous_area_emergency_type[]')
            },
            'terror_consequences': {
                'num': request.form.getlist('terror_consequence_num[]'),
                'threat': request.form.getlist('terror_consequence_threat[]'),
                'victims_count': request.form.getlist('terror_consequence_victims_count[]'),
                'consequence_scale': request.form.getlist('terror_consequence_consequence_scale[]')
            },
            'security_posts': {
                'post_type': request.form.getlist('security_post_type[]'),
                'units': request.form.getlist('security_post_units[]'),
                'persons': request.form.getlist('security_post_persons[]')
            },
            'protection_assessment': {
                'num': request.form.getlist('protection_assessment_num[]'),
                'element_name': request.form.getlist('protection_assessment_element_name[]'),
                'requirements': request.form.getlist('protection_assessment_requirements[]'),
                'physical_protection': request.form.getlist('protection_assessment_physical_protection[]'),
                'terror_prevention': request.form.getlist('protection_assessment_terror_prevention[]'),
                'sufficiency': request.form.getlist('protection_assessment_sufficiency[]'),
                'compensation': request.form.getlist('protection_assessment_compensation[]')
            }
        }

        # Загружаем шаблон
        doc = Document(template_path)
        logger.info("Шаблон успешно загружен")

        # Функция для замены плейсхолдеров в тексте
        def replace_text(text, fields, table_fields):
            if not text:
                return text
            
            # Замена простых полей
            for key, placeholder in text_fields.items():
                if placeholder in text:
                    value = fields.get(key, '').strip()
                    text = text.replace(placeholder, value if value else "")
            
            # Замена плейсхолдеров для таблиц (кроме пунктов 4 и 10г)
            for table_key in [k for k in table_fields.keys() if k not in ['transport_communications', 'security_posts']]:
                table_data = table_fields[table_key]
                row_count = max(len(table_data[key]) for key in table_data if table_data[key]) if any(table_data[key] for key in table_data) else 0
                for i in range(row_count):
                    for sub_key in table_data.keys():
                        placeholder = f"{{{table_key}[{i}].{sub_key}}}"
                        if placeholder in text:
                            value = table_data[sub_key][i] if i < len(table_data[sub_key]) and table_data[sub_key][i] else ""
                            text = text.replace(placeholder, str(value))
            return text

        # Функция для замены плейсхолдеров в таблицах
        def replace_table_placeholders(table, fields, table_fields):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        new_text = replace_text(original_text, fields, table_fields)
                        if new_text != original_text:
                            para.text = new_text
                            logger.info(f"Замена в таблице: '{original_text}' -> '{new_text}'")

        # Замена плейсхолдеров в параграфах
        for para in doc.paragraphs:
            original_text = para.text
            new_text = replace_text(original_text, fields, table_fields)
            if new_text != original_text:
                para.text = new_text
                para.alignment = 1  # Выравнивание по центру
                logger.info(f"Замена в параграфе: '{original_text}' -> '{new_text}'")

        # Замена плейсхолдеров в таблицах
        for table in doc.tables:
            replace_table_placeholders(table, fields, table_fields)

        # Сохраняем документ
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        logger.info("Документ успешно сохранен")

        return send_file(output, download_name="Паспорт_безопасности.docx", as_attachment=True)

    except Exception as e:
        logger.error(f"Ошибка при генерации документа: {str(e)}")
        return f"Ошибка при генерации документа: {str(e)}", 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
